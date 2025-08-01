import os
import json
import csv
from pathlib import Path
from typing import Dict, List, Any, Optional
import mimetypes

# Import library untuk berbagai format dokumen
try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    print("Install python-docx: pip install python-docx")

try:
    import openpyxl
    from openpyxl import Workbook
except ImportError:
    print("Install openpyxl: pip install openpyxl")

try:
    from PyPDF2 import PdfReader, PdfWriter
except ImportError:
    print("Install PyPDF2: pip install PyPDF2")

try:
    import xml.etree.ElementTree as ET
except ImportError:
    pass

class DocumentProcessor:
    """Kelas utama untuk memproses berbagai jenis dokumen"""
    
    def __init__(self):
        self.supported_formats = {
            '.txt': self._process_text,
            '.docx': self._process_docx,
            '.xlsx': self._process_excel,
            '.csv': self._process_csv,
            '.json': self._process_json,
            '.xml': self._process_xml,
            '.pdf': self._process_pdf
        }
    
    def detect_document_type(self, file_path: str) -> str:
        """Mendeteksi tipe dokumen berdasarkan ekstensi dan MIME type"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File {file_path} tidak ditemukan")
        
        # Deteksi berdasarkan ekstensi
        file_extension = Path(file_path).suffix.lower()
        
        # Deteksi berdasarkan MIME type sebagai backup
        mime_type, _ = mimetypes.guess_type(file_path)
        
        print(f"File: {file_path}")
        print(f"Ekstensi: {file_extension}")
        print(f"MIME Type: {mime_type}")
        
        return file_extension
    
    def process_document(self, file_path: str, operation: str, **kwargs) -> Any:
        """Memproses dokumen berdasarkan tipe dan operasi yang diminta"""
        doc_type = self.detect_document_type(file_path)
        
        if doc_type not in self.supported_formats:
            raise ValueError(f"Format {doc_type} tidak didukung. Format yang didukung: {list(self.supported_formats.keys())}")
        
        processor = self.supported_formats[doc_type]
        return processor(file_path, operation, **kwargs)
    
    def _process_text(self, file_path: str, operation: str, **kwargs) -> Any:
        """Memproses file teks (.txt)"""
        if operation == "read":
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        
        elif operation == "write":
            content = kwargs.get('content', '')
            with open(file_path, 'w', encoding='utf-8') as file:
                file.write(content)
            return f"File {file_path} berhasil ditulis"
        
        elif operation == "append":
            content = kwargs.get('content', '')
            with open(file_path, 'a', encoding='utf-8') as file:
                file.write('\n' + content)
            return f"Konten berhasil ditambahkan ke {file_path}"
        
        elif operation == "replace":
            old_text = kwargs.get('old_text', '')
            new_text = kwargs.get('new_text', '')
            
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            updated_content = content.replace(old_text, new_text)
            
            with open(file_path, 'w', encoding='utf-8') as file:
                file.write(updated_content)
            
            return f"Teks '{old_text}' berhasil diganti dengan '{new_text}'"
    
    def _process_docx(self, file_path: str, operation: str, **kwargs) -> Any:
        """Memproses file Word (.docx)"""
        if operation == "read":
            doc = Document(file_path)
            text_content = []
            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)
            return '\n'.join(text_content)
        
        elif operation == "write":
            content = kwargs.get('content', '')
            doc = Document()
            doc.add_paragraph(content)
            doc.save(file_path)
            return f"Dokumen Word {file_path} berhasil dibuat"
        
        elif operation == "append":
            content = kwargs.get('content', '')
            doc = Document(file_path)
            doc.add_paragraph(content)
            doc.save(file_path)
            return f"Paragraf baru berhasil ditambahkan ke {file_path}"
        
        elif operation == "replace":
            old_text = kwargs.get('old_text', '')
            new_text = kwargs.get('new_text', '')
            
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                if old_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_text, new_text)
            
            doc.save(file_path)
            return f"Teks dalam dokumen Word berhasil diganti"
    
    def _process_excel(self, file_path: str, operation: str, **kwargs) -> Any:
        """Memproses file Excel (.xlsx)"""
        if operation == "read":
            workbook = openpyxl.load_workbook(file_path)
            sheet = workbook.active
            data = []
            for row in sheet.iter_rows(values_only=True):
                data.append(row)
            return data
        
        elif operation == "write":
            data = kwargs.get('data', [])
            workbook = Workbook()
            sheet = workbook.active
            
            for row_data in data:
                sheet.append(row_data)
            
            workbook.save(file_path)
            return f"File Excel {file_path} berhasil dibuat"
        
        elif operation == "append":
            new_row = kwargs.get('row_data', [])
            workbook = openpyxl.load_workbook(file_path)
            sheet = workbook.active
            sheet.append(new_row)
            workbook.save(file_path)
            return f"Baris baru berhasil ditambahkan ke {file_path}"
        
        elif operation == "update_cell":
            row = kwargs.get('row', 1)
            col = kwargs.get('col', 1)
            value = kwargs.get('value', '')
            
            workbook = openpyxl.load_workbook(file_path)
            sheet = workbook.active
            sheet.cell(row=row, column=col, value=value)
            workbook.save(file_path)
            return f"Cell ({row}, {col}) berhasil diupdate"
    
    def _process_csv(self, file_path: str, operation: str, **kwargs) -> Any:
        """Memproses file CSV"""
        if operation == "read":
            data = []
            with open(file_path, 'r', encoding='utf-8') as file:
                csv_reader = csv.reader(file)
                for row in csv_reader:
                    data.append(row)
            return data
        
        elif operation == "write":
            data = kwargs.get('data', [])
            with open(file_path, 'w', newline='', encoding='utf-8') as file:
                csv_writer = csv.writer(file)
                csv_writer.writerows(data)
            return f"File CSV {file_path} berhasil dibuat"
        
        elif operation == "append":
            new_row = kwargs.get('row_data', [])
            with open(file_path, 'a', newline='', encoding='utf-8') as file:
                csv_writer = csv.writer(file)
                csv_writer.writerow(new_row)
            return f"Baris baru berhasil ditambahkan ke CSV"
    
    def _process_json(self, file_path: str, operation: str, **kwargs) -> Any:
        """Memproses file JSON"""
        if operation == "read":
            with open(file_path, 'r', encoding='utf-8') as file:
                return json.load(file)
        
        elif operation == "write":
            data = kwargs.get('data', {})
            with open(file_path, 'w', encoding='utf-8') as file:
                json.dump(data, file, indent=2, ensure_ascii=False)
            return f"File JSON {file_path} berhasil dibuat"
        
        elif operation == "update":
            key = kwargs.get('key', '')
            value = kwargs.get('value', '')
            
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            # Update nested keys dengan dot notation
            keys = key.split('.')
            current = data
            for k in keys[:-1]:
                if k not in current:
                    current[k] = {}
                current = current[k]
            current[keys[-1]] = value
            
            with open(file_path, 'w', encoding='utf-8') as file:
                json.dump(data, file, indent=2, ensure_ascii=False)
            
            return f"Key '{key}' berhasil diupdate"
    
    def _process_xml(self, file_path: str, operation: str, **kwargs) -> Any:
        """Memproses file XML"""
        if operation == "read":
            tree = ET.parse(file_path)
            root = tree.getroot()
            return ET.tostring(root, encoding='unicode')
        
        elif operation == "write":
            root_name = kwargs.get('root_name', 'root')
            data = kwargs.get('data', {})
            
            root = ET.Element(root_name)
            self._dict_to_xml(data, root)
            
            tree = ET.ElementTree(root)
            tree.write(file_path, encoding='utf-8', xml_declaration=True)
            return f"File XML {file_path} berhasil dibuat"
        
        elif operation == "add_element":
            parent_xpath = kwargs.get('parent_xpath', '.')
            element_name = kwargs.get('element_name', 'new_element')
            element_text = kwargs.get('element_text', '')
            
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            parent = root.find(parent_xpath)
            if parent is not None:
                new_element = ET.SubElement(parent, element_name)
                new_element.text = element_text
                
                tree.write(file_path, encoding='utf-8', xml_declaration=True)
                return f"Element '{element_name}' berhasil ditambahkan"
            else:
                return f"Parent element tidak ditemukan: {parent_xpath}"
    
    def _process_pdf(self, file_path: str, operation: str, **kwargs) -> Any:
        """Memproses file PDF (hanya baca untuk sekarang)"""
        if operation == "read":
            reader = PdfReader(file_path)
            text_content = []
            for page in reader.pages:
                text_content.append(page.extract_text())
            return '\n'.join(text_content)
        
        else:
            return "Operasi PDF terbatas pada pembacaan saja"
    
    def _dict_to_xml(self, data: Dict, parent: ET.Element):
        """Helper untuk mengkonversi dictionary ke XML"""
        for key, value in data.items():
            if isinstance(value, dict):
                child = ET.SubElement(parent, key)
                self._dict_to_xml(value, child)
            else:
                child = ET.SubElement(parent, key)
                child.text = str(value)


# Contoh penggunaan aplikasi
def main():
    """Fungsi utama untuk demo aplikasi"""
    processor = DocumentProcessor()
    
    print("=== APLIKASI MANIPULASI DOKUMEN ===\n")
    
    # Contoh penggunaan untuk berbagai format
    try:
        # 1. File Teks
        print("1. DEMO FILE TEKS")
        processor.process_document("test.txt", "write", content="Halo, ini adalah file teks pertama.")
        content = processor.process_document("test.txt", "read")
        print(f"Isi file: {content}")
        processor.process_document("test.txt", "append", content="Baris tambahan.")
        processor.process_document("test.txt", "replace", old_text="pertama", new_text="yang telah dimodifikasi")
        print()
        
        # 2. File CSV
        print("2. DEMO FILE CSV")
        csv_data = [
            ["Nama", "Umur", "Kota"],
            ["Alice", 25, "Jakarta"],
            ["Bob", 30, "Bandung"]
        ]
        processor.process_document("test.csv", "write", data=csv_data)
        processor.process_document("test.csv", "append", row_data=["Charlie", 28, "Surabaya"])
        csv_content = processor.process_document("test.csv", "read")
        print(f"Data CSV: {csv_content}")
        print()
        
        # 3. File JSON
        print("3. DEMO FILE JSON")
        json_data = {
            "aplikasi": "Document Processor",
            "versi": "1.0",
            "fitur": ["read", "write", "update", "delete"]
        }
        processor.process_document("test.json", "write", data=json_data)
        processor.process_document("test.json", "update", key="versi", value="1.1")
        processor.process_document("test.json", "update", key="author.nama", value="Developer")
        json_content = processor.process_document("test.json", "read")
        print(f"Data JSON: {json.dumps(json_content, indent=2)}")
        print()
        
        # 4. File Excel (jika library tersedia)
        print("4. DEMO FILE EXCEL")
        excel_data = [
            ["Produk", "Harga", "Stok"],
            ["Laptop", 10000000, 5],
            ["Mouse", 50000, 20]
        ]
        try:
            processor.process_document("test.xlsx", "write", data=excel_data)
            processor.process_document("test.xlsx", "append", row_data=["Keyboard", 150000, 15])
            processor.process_document("test.xlsx", "update_cell", row=2, col=3, value=8)
            print("File Excel berhasil dibuat dan dimanipulasi")
        except Exception as e:
            print(f"Error Excel: {e}")
        print()
        
        print("Demo selesai! File-file test telah dibuat.")
        
    except Exception as e:
        print(f"Error: {e}")


if __name__ == "__main__":
    main()
